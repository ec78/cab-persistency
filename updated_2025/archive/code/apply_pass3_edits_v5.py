"""
apply_pass3_edits_v5.py
=======================
Pass 3 edits: working_draft_ec_2025_v4_Mar2026.docx → working_draft_ec_2025_v5_Mar2026.docx

Addresses remaining gaps from paper_review.md and DOCUMENTATION_PERSONA_FEEDBACK.md:
  P1  Narrow framing in intro (para 16): clarify step 2 explains explosive vs non-explosive
  P2  Clower & Ito differentiation: new paragraph after intro roadmap
  P3  Rare events limitation: new paragraph after IIA paragraph
  P4  Soften causal language in AME discussion (paras 137, 163, 169)
  P5  Explosive recall framing: clarify significance ≠ predictive accuracy (para 139)
  P6  Post-GFC tone-down: hedge institutional interpretation (para 161)
  P7  Annual aggregation defense: add sentence in Section 3.2 (para 62)
  P8  Abstract typo: "in the the full sample" → "in the full sample"

Run from the repo root or this script's directory. Requires: python-docx.
"""

import copy
import sys
from docx import Document
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

INPUT  = r"c:\Users\eclow\Documents\GitHub\cab-persistency\updated_2025\working_draft_ec_2025_v4_Mar2026.docx"
OUTPUT = r"c:\Users\eclow\Documents\GitHub\cab-persistency\updated_2025\working_draft_ec_2025_v5_Mar2026.docx"


# ---------------------------------------------------------------------------
# Helper utilities
# ---------------------------------------------------------------------------

def para_text(p):
    return p.text.strip()


def replace_text_in_para(para, old, new):
    """Replace old with new across all runs in a paragraph, handling cross-run splits."""
    # Simple case: whole text replacement via full-text match
    full = para.text
    if old not in full:
        return False
    # Clear existing runs and set new text as single run (preserves first run's style)
    if para.runs:
        para.runs[0].text = full.replace(old, new)
        for run in para.runs[1:]:
            run.text = ""
    else:
        para.add_run(full.replace(old, new))
    return True


def insert_paragraph_after(reference_para, text, style=None):
    """Insert a new paragraph with text immediately after reference_para."""
    new_para = OxmlElement('w:p')
    # Copy paragraph properties (style) from reference if present
    ref_pPr = reference_para._element.find(qn('w:pPr'))
    if ref_pPr is not None:
        new_pPr = copy.deepcopy(ref_pPr)
        new_para.append(new_pPr)
    # Add run with text
    new_r = OxmlElement('w:r')
    # Copy run properties from first run of reference paragraph if available
    if reference_para.runs:
        ref_rPr = reference_para.runs[0]._element.find(qn('w:rPr'))
        if ref_rPr is not None:
            new_rPr = copy.deepcopy(ref_rPr)
            new_r.append(new_rPr)
    new_t = OxmlElement('w:t')
    new_t.text = text
    new_t.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')
    new_r.append(new_t)
    new_para.append(new_r)
    # Insert after reference_para in the document body
    reference_para._element.addnext(new_para)
    return new_para


def find_para_by_prefix(doc, prefix, skip=0):
    """Return the paragraph whose text starts with prefix (case-sensitive).
    skip=n skips the first n matches."""
    count = 0
    for p in doc.paragraphs:
        if para_text(p).startswith(prefix):
            if count == skip:
                return p
            count += 1
    return None


def find_para_containing(doc, substring, skip=0):
    """Return the paragraph whose text contains substring."""
    count = 0
    for p in doc.paragraphs:
        if substring in para_text(p):
            if count == skip:
                return p
            count += 1
    return None


# ---------------------------------------------------------------------------
# Load document
# ---------------------------------------------------------------------------

doc = Document(INPUT)

# Build list of non-empty paragraphs for index reference (matches earlier analysis)
paras = [p for p in doc.paragraphs if p.text.strip()]
print(f"Loaded document. Non-empty paragraphs: {len(paras)}")


# ---------------------------------------------------------------------------
# P8: Abstract typo — "in the the full sample" → "in the full sample"
# ---------------------------------------------------------------------------

abstract_para = find_para_containing(doc, "in the the full sample")
if abstract_para:
    result = replace_text_in_para(abstract_para, "in the the full sample", "in the full sample")
    print(f"P8 (abstract typo): {'done' if result else 'FAILED'}")
else:
    print("P8 (abstract typo): paragraph not found")


# ---------------------------------------------------------------------------
# P1: Narrow framing in intro para 16
# Current: "Our analysis focuses on persistence and the structural factors that govern it."
# New: also add a sentence clarifying that step 2 explains explosive vs non-explosive.
# We insert the clarifying sentence at the end of the paragraph.
# ---------------------------------------------------------------------------

PARA16_PREFIX = "Against this backdrop, this paper examines"
p16 = find_para_by_prefix(doc, PARA16_PREFIX)
if p16:
    OLD16 = (
        "Finally, we examine whether the likelihood of entering different persistence regimes"
        "\u2014and the degree of persistence within those regimes\u2014can be explained by "
        "cross-sectional and temporal variation in macroeconomic policies, institutional "
        "characteristics, and structural fundamentals."
    )
    NEW16 = (
        "Finally, we examine whether the likelihood of entering different persistence regimes "
        "\u2014and in particular the explosive regime, where external dynamics become "
        "destabilizing\u2014can be explained by cross-sectional and temporal variation in "
        "macroeconomic policies, institutional characteristics, and structural fundamentals. "
        "We show that structural factors are most informative for explaining which countries "
        "enter explosive adjustment paths; within the stable regime space, the same variables "
        "do not reliably distinguish stationary from unit-root behavior."
    )
    result = replace_text_in_para(p16, OLD16, NEW16)
    if not result:
        # Try a shorter unique substring
        OLD16_SHORT = "Finally, we examine whether the likelihood of entering different persistence regimes"
        full = p16.text
        if OLD16_SHORT in full:
            new_full = full.replace(
                "Finally, we examine whether the likelihood of entering different persistence regimes"
                "\u2014and the degree of persistence within those regimes\u2014can be explained by "
                "cross-sectional and temporal variation in macroeconomic policies, institutional "
                "characteristics, and structural fundamentals.",
                NEW16
            )
            if p16.runs:
                p16.runs[0].text = new_full
                for r in p16.runs[1:]:
                    r.text = ""
            result = True
    print(f"P1 (intro framing): {'done' if result else 'FAILED - text not found'}")
    if not result:
        print(f"  Para text: {p16.text[:200]}")
else:
    print("P1 (intro framing): paragraph not found")


# ---------------------------------------------------------------------------
# P2: Clower & Ito differentiation — insert new paragraph after para 16
# ---------------------------------------------------------------------------

if p16:
    CLOWER_ITO_TEXT = (
        "Our paper extends a small but growing literature that applies Markov-switching "
        "models to current account dynamics. Earlier work in this lineage studied current "
        "account persistence using two-state Markov-switching classifications and identified "
        "structural predictors of entry into a random-walk regime. Our contribution relative "
        "to that work is threefold. First, we adopt the independent-switching specification "
        "of Morita et al. (2024), which allows persistence, mean, and volatility to switch "
        "separately across regimes; this prevents changes in mean or volatility from being "
        "spuriously attributed to persistence shifts, a problem that biases classification in "
        "standard two-state models. Second, we extend to a three-regime framework that "
        "isolates an explosive tail separately from persistent-but-bounded dynamics. The "
        "two-state (stationary vs. random walk) framework cannot produce the central result "
        "of this paper \u2014 that structural variables matter primarily for explosive regime "
        "entry rather than for the stationary-unit-root distinction. Third, we add a formal "
        "multinomial logit second stage with cross-validated model selection and multiple "
        "robustness checks, including a binary logit that precisely bounds what the model "
        "can and cannot explain."
    )
    insert_paragraph_after(p16, CLOWER_ITO_TEXT)
    print("P2 (Clower & Ito differentiation paragraph): inserted after intro roadmap para")
else:
    print("P2: reference paragraph not found; skipped")


# ---------------------------------------------------------------------------
# P3: Rare events limitation — insert new paragraph after IIA paragraph (para 127)
# ---------------------------------------------------------------------------

IIA_PREFIX = "Second, the multinomial logit rests on the independence of irrelevant alternatives"
p_iia = find_para_by_prefix(doc, IIA_PREFIX)
if p_iia:
    RARE_EVENTS_TEXT = (
        "A third limitation concerns inference in the presence of rare events. The explosive "
        "regime comprises only 5.5 percent of the sample \u2014 134 country-years concentrated "
        "in a small number of economies (principally Israel, Hungary, Malaysia, Germany, and "
        "Armenia). All MNL coefficient estimates for the explosive equation rely on contrasts "
        "with this geographically concentrated reference group, raising concerns about "
        "small-sample bias in the explosive-equation parameters and sensitivity to influential "
        "observations. A fully rigorous treatment would include leave-one-country-out and "
        "leave-one-region-out sensitivity analysis and a Firth penalized logistic regression "
        "for the binary explosive-versus-non-explosive contrast; these extensions are left for "
        "future work. The lagged exchange rate robustness (Section 5.3.3) and the binary logit "
        "exercise (Section 5.4) provide partial assurance that the headline finding is not "
        "driven by a single influential country, but they do not constitute a full rare-events "
        "robustness analysis."
    )
    insert_paragraph_after(p_iia, RARE_EVENTS_TEXT)
    print("P3 (rare events limitation paragraph): inserted after IIA paragraph")
else:
    print("P3: IIA paragraph not found; skipped")
    # Try alternative search
    p_iia_alt = find_para_containing(doc, "independence of irrelevant alternatives")
    if p_iia_alt:
        insert_paragraph_after(p_iia_alt, RARE_EVENTS_TEXT)
        print("P3 (rare events): inserted using alternative search")


# ---------------------------------------------------------------------------
# P4: Soften causal language in AME discussion
# ---------------------------------------------------------------------------

# Para 137: "Fixed exchange rate arrangements reduce the probability of the explosive regime"
AME_PREFIX = "The marginal effects reveal a clear asymmetry"
p_ame = find_para_by_prefix(doc, AME_PREFIX)
if p_ame:
    result = replace_text_in_para(
        p_ame,
        "Fixed exchange rate arrangements reduce the probability of the explosive regime by approximately 9.1 percentage points.",
        "Fixed exchange rate arrangements are associated with a reduction in the probability of the explosive regime of approximately 9.1 percentage points."
    )
    print(f"P4a (AME causal softening): {'done' if result else 'FAILED'}")
    if not result:
        print(f"  Para text snippet: {p_ame.text[:300]}")
else:
    print("P4a: AME paragraph not found")

# Para 163 (Section 5.6 summary): "reduces the probability of the explosive regime"
SUMMARY_PREFIX = "The multinomial logit analysis yields three principal findings. First"
p_sum = find_para_by_prefix(doc, SUMMARY_PREFIX)
if p_sum:
    result = replace_text_in_para(
        p_sum,
        "Countries under fixed exchange rate arrangements are significantly and substantially less likely to enter the explosive regime, with average marginal effects indicating an approximately 9 percentage point reduction in the probability of explosive behavior.",
        "Countries under fixed exchange rate arrangements are significantly and substantially less likely to enter the explosive regime, with average marginal effects indicating an approximately 9 percentage point association with lower probability of explosive behavior."
    )
    print(f"P4b (summary causal softening): {'done' if result else 'FAILED'}")
else:
    print("P4b: summary paragraph not found")

# Para 169 (conclusion): "Fixed pegs substantially reduce the probability"
CONC_PREFIX = "Second, we identify fixed exchange rate arrangements"
p_conc = find_para_by_prefix(doc, CONC_PREFIX)
if p_conc:
    result = replace_text_in_para(
        p_conc,
        "Fixed pegs substantially reduce the probability of explosive current account dynamics, consistent with the theoretical prediction that exchange rate commitments constrain the scope for unchecked external imbalances by limiting expenditure-switching.",
        "Fixed pegs are associated with substantially lower probability of explosive current account dynamics, consistent with the theoretical prediction that exchange rate commitments constrain the scope for unchecked external imbalances by limiting expenditure-switching."
    )
    print(f"P4c (conclusion causal softening): {'done' if result else 'FAILED'}")
else:
    print("P4c: conclusion paragraph not found")


# ---------------------------------------------------------------------------
# P5: Explosive recall framing — add clarifying sentence to para 139
# ---------------------------------------------------------------------------

CONFMAT_PREFIX = "Table C1 in the Appendix reports a confusion matrix"
p_confmat = find_para_by_prefix(doc, CONFMAT_PREFIX)
if p_confmat:
    OLD_CONFMAT = "Country-grouped five-fold cross-validation yields a mean log-loss of 0.945, indicating modest but meaningful out-of-sample predictive content."
    NEW_CONFMAT = (
        "Country-grouped five-fold cross-validation yields a mean log-loss of 0.945, "
        "indicating modest but meaningful out-of-sample predictive content. "
        "It is important to distinguish coefficient significance from predictive performance: "
        "the statistical significance of the exchange rate regime coefficient reflects a "
        "systematic shift in the probability distribution of regime outcomes, not a reliable "
        "ability to flag individual explosive episodes. A variable can strongly shift the "
        "probability of a rare event without enabling accurate event-by-event prediction, "
        "and the confusion matrix confirms that this is the case here."
    )
    result = replace_text_in_para(p_confmat, OLD_CONFMAT, NEW_CONFMAT)
    print(f"P5 (explosive recall framing): {'done' if result else 'FAILED'}")
    if not result:
        print(f"  Para text snippet: {p_confmat.text[:300]}")
else:
    print("P5: confusion matrix paragraph not found")


# ---------------------------------------------------------------------------
# P6: Post-GFC tone-down — hedge institutional interpretation (para 161)
# ---------------------------------------------------------------------------

POSTGFC_PREFIX = "Model fit more than doubles in the post-GFC period"
p_gfc = find_para_by_prefix(doc, POSTGFC_PREFIX)
if p_gfc:
    OLD_GFC1 = "Three institutional developments reinforce this interpretation:"
    NEW_GFC1 = "Three institutional developments are plausibly consistent with this interpretation, though they are not directly measured in the data:"
    r1 = replace_text_in_para(p_gfc, OLD_GFC1, NEW_GFC1)

    OLD_GFC2 = (
        "Crucially, this improvement in fit is not explained by a common post-crisis time trend: "
        "the targeted crisis-dummy robustness check confirms that global shocks add no predictive "
        "value once structural covariates are controlled. Rather, the post-GFC improvement reflects "
        "a structural change in how cross-country differences in policy frameworks translate into "
        "external regime behavior."
    )
    NEW_GFC2 = (
        "Crucially, this improvement in fit is not explained by a common post-crisis time trend: "
        "the targeted crisis-dummy robustness check confirms that global shocks add no predictive "
        "value once structural covariates are controlled. One interpretation \u2014 albeit not "
        "directly tested, since these institutional developments are not measured in the data "
        "\u2014 is that the post-GFC period represents a structural change in how cross-country "
        "differences in policy frameworks translate into external regime behavior."
    )
    r2 = replace_text_in_para(p_gfc, OLD_GFC2, NEW_GFC2)
    print(f"P6 (post-GFC hedge): part1={'done' if r1 else 'FAILED'}, part2={'done' if r2 else 'FAILED'}")
    if not r1:
        print(f"  Para text snippet: {p_gfc.text[:400]}")
else:
    print("P6: post-GFC paragraph not found")


# ---------------------------------------------------------------------------
# P7: Annual aggregation defense — add sentence to Section 3.2 para (para 62)
# ---------------------------------------------------------------------------

AR1_PREFIX = "Formally, let denote the current account balance as a percentage of GDP"
p_ar1 = find_para_by_prefix(doc, AR1_PREFIX)
if p_ar1:
    OLD_AR1 = (
        "To ensure comparability across countries, we impose a common AR(1) structure for all series. "
        "Allowing country-specific lag orders would complicate cross-country aggregation and is therefore left for future work."
    )
    NEW_AR1 = (
        "To ensure comparability across countries, we impose a common AR(1) structure for all series. "
        "Allowing country-specific lag orders would complicate cross-country aggregation and is therefore left for future work. "
        "For the second stage, quarterly regime classifications are aggregated to the annual level by "
        "taking the modal (most frequent) regime within each country-year. This aggregation matches "
        "the annual frequency of the structural covariates and produces a single interpretable outcome "
        "per observation. Within-year transitions are collapsed to the modal state, discarding "
        "potentially informative intra-year dynamics; more granular temporal aggregation approaches "
        "are left for future work."
    )
    result = replace_text_in_para(p_ar1, OLD_AR1, NEW_AR1)
    print(f"P7 (annual aggregation defense): {'done' if result else 'FAILED'}")
    if not result:
        print(f"  Para text snippet: {p_ar1.text[:300]}")
else:
    print("P7: AR(1) paragraph not found")


# ---------------------------------------------------------------------------
# Save output
# ---------------------------------------------------------------------------

doc.save(OUTPUT)
print(f"\nSaved: {OUTPUT}")
print("Pass 3 complete.")
