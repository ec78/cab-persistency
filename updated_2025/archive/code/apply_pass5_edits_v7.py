"""
Pass 5 paper edits — working_draft_ec_2025_v7_Mar2026.docx
Source: date-range correction (1985→1971) + unbalanced coverage note.

Changes:
  E1 — Abstract/intro sample period: "1985–2023" → "1971–2023" (para 8 approx)
  E2 — Section 5.1.2 sample description: "1985–2023" → "1971–2023" + add coverage sentence (para 150 approx)
  E3 — Conclusion sample period: "1985–2023" → "1971–2023" + brief parenthetical (para 204 approx)

Run from updated_2025/:
    python "code/mnl/python code/apply_pass5_edits_v7.py"
"""

import re
from docx import Document
from docx.oxml.ns import qn
import copy

INPUT_DOCX  = "working_draft_ec_2025_v6_Mar2026.docx"
OUTPUT_DOCX = "working_draft_ec_2025_v7_Mar2026.docx"

# ---------------------------------------------------------------------------
# Helpers
# ---------------------------------------------------------------------------

def para_text(para):
    """Return full text of a paragraph (joining all runs)."""
    return "".join(r.text for r in para.runs)


def replace_in_para(para, old, new):
    """
    Replace the first occurrence of `old` with `new` across the paragraph's
    runs, preserving run formatting.  Returns True if a replacement was made.
    """
    full = para_text(para)
    if old not in full:
        return False

    # Rebuild runs: collapse to a single text replacement on the first match,
    # keeping all runs' styles from the first run that touches the match.
    new_full = full.replace(old, new, 1)

    # Simplest safe approach: dump everything into the first run, clear others.
    if para.runs:
        para.runs[0].text = new_full
        for r in para.runs[1:]:
            r.text = ""
    return True


def append_sentence_to_para(para, sentence):
    """
    Append `sentence` to the last run of `para`, or create a new run if none
    exist.  Inherits the formatting of the last run.
    """
    if para.runs:
        last_run = para.runs[-1]
        last_run.text = last_run.text.rstrip() + " " + sentence
    else:
        run = para.add_run(" " + sentence)


def insert_para_after(ref_para, text, style=None):
    """Insert a new paragraph with `text` immediately after `ref_para`."""
    new_para = copy.deepcopy(ref_para)
    # Clear all runs in the copy
    for child in list(new_para._element):
        if child.tag.endswith("}r"):
            new_para._element.remove(child)
    # Add a single run with the text
    from docx.oxml import OxmlElement
    r_elem = OxmlElement("w:r")
    rpr = ref_para.runs[0]._r.find(qn("w:rPr")) if ref_para.runs else None
    if rpr is not None:
        r_elem.append(copy.deepcopy(rpr))
    t_elem = OxmlElement("w:t")
    t_elem.text = text
    t_elem.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
    r_elem.append(t_elem)
    new_para._element.append(r_elem)
    ref_para._element.addnext(new_para._element)
    return new_para


# ---------------------------------------------------------------------------
# Load document
# ---------------------------------------------------------------------------
doc = Document(INPUT_DOCX)
paras = doc.paragraphs

# ---------------------------------------------------------------------------
# E1 — Abstract / intro: update sample period statement
#       Target: "74 countries over the period 1985–2023"
# ---------------------------------------------------------------------------
E1_OLD = "74 countries over the period 1985\u20132023"
E1_NEW = "74 countries spanning 1971\u20132023"

e1_done = False
for i, para in enumerate(paras):
    t = para_text(para)
    if E1_OLD in t:
        replace_in_para(para, E1_OLD, E1_NEW)
        print(f"[E1] Applied at paragraph {i}: {t[:80]}...")
        e1_done = True
        break

if not e1_done:
    print("[E1] WARNING: target text not found — check para text for E1.")

# ---------------------------------------------------------------------------
# E2 — Section 5.1.2 sample description
#       Target: "The complete sample covers 74 countries observed annually
#                over the period 1985–2023, yielding 2,451 usable observations
#                after listwise deletion."
#       Action: update period AND append a coverage sentence.
# ---------------------------------------------------------------------------
E2_OLD_PERIOD = "over the period 1985\u20132023"
E2_NEW_PERIOD = "spanning 1971\u20132023"
E2_ANCHOR = "2,451 usable observations after listwise deletion"
E2_COVERAGE_SENTENCE = (
    "Country coverage is strongly unbalanced in the early part of the sample: "
    "fewer than 30 countries contribute observations before 1985, "
    "rising to 39 countries by 1995 and 70 or more from 2008 onward. "
    "The analysis retains these early observations rather than imposing a "
    "common start date, as the unbalanced structure reflects genuine data "
    "availability rather than sample selection."
)

e2_done = False
for i, para in enumerate(paras):
    t = para_text(para)
    if E2_ANCHOR in t:
        # First update the period phrase
        if E2_OLD_PERIOD in t:
            replace_in_para(para, E2_OLD_PERIOD, E2_NEW_PERIOD)
            print(f"[E2a] Period updated at paragraph {i}.")
        else:
            print(f"[E2a] Period phrase not found in anchor paragraph {i}; skipping period update.")
        # Now append the coverage sentence
        append_sentence_to_para(para, E2_COVERAGE_SENTENCE)
        print(f"[E2b] Coverage sentence appended at paragraph {i}.")
        e2_done = True
        break

if not e2_done:
    print("[E2] WARNING: anchor text not found — check para text for E2.")

# ---------------------------------------------------------------------------
# E3 — Conclusion: update sample period
#       Target: "74 countries using annual data spanning 1985–2023"
# ---------------------------------------------------------------------------
E3_OLD = "annual data spanning 1985\u20132023"
E3_NEW = ("annual data spanning 1971\u20132023 "
          "(with country coverage growing from a small early-adopting group "
          "to 70 or more countries from 2008 onward)")

e3_done = False
for i, para in enumerate(paras):
    t = para_text(para)
    if E3_OLD in t:
        replace_in_para(para, E3_OLD, E3_NEW)
        print(f"[E3] Applied at paragraph {i}: {t[:80]}...")
        e3_done = True
        break

if not e3_done:
    print("[E3] WARNING: target text not found — check para text for E3.")

# ---------------------------------------------------------------------------
# Save
# ---------------------------------------------------------------------------
doc.save(OUTPUT_DOCX)
print(f"\nSaved: {OUTPUT_DOCX}")
print(f"  E1 done: {e1_done}")
print(f"  E2 done: {e2_done}")
print(f"  E3 done: {e3_done}")
