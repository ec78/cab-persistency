"""
apply_pass2_edits_v2.py
Corrected version — applies all Pass 2 edits from v3 docx with better paragraph targeting.
Saves as working_draft_ec_2025_v4_Mar2026.docx.
"""

import sys, io
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding="utf-8", errors="replace")

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
import copy

IN  = (r"c:/Users/eclow/Documents/GitHub/cab-persistency/"
       r"updated_2025/working_draft_ec_2025_v3_Mar2026.docx")
OUT = (r"c:/Users/eclow/Documents/GitHub/cab-persistency/"
       r"updated_2025/working_draft_ec_2025_v4_Mar2026.docx")


# ---------------------------------------------------------------------------
# Helpers
# ---------------------------------------------------------------------------

def find_para_idx(doc, search_text, start=0, end=None, case=False):
    """Return first paragraph index (>= start, < end) containing search_text."""
    needle = search_text if case else search_text.lower()
    paras = doc.paragraphs
    if end is None:
        end = len(paras)
    for i in range(start, min(end, len(paras))):
        hay = paras[i].text if case else paras[i].text.lower()
        if needle in hay:
            return i
    return None


def clone_para_format(source):
    """Return a deep-copied empty <w:p> element matching source's paragraph properties."""
    new_elem = copy.deepcopy(source._element)
    for r in new_elem.findall(qn("w:r")):
        new_elem.remove(r)
    return new_elem


def insert_paragraph_after(ref_para, text, bold=False, italic=False):
    """Insert a new paragraph after ref_para with the given text. Returns new Paragraph."""
    from docx.text.paragraph import Paragraph as DocxPara
    new_elem = clone_para_format(ref_para)
    ref_para._element.addnext(new_elem)
    new_p = DocxPara(new_elem, ref_para._parent)
    run = new_p.add_run(text)
    run.bold   = bold
    run.italic = italic
    if ref_para.runs:
        src = ref_para.runs[0].font
        f   = run.font
        f.name, f.size, f.bold, f.italic = src.name, src.size, src.bold, src.italic
        run.bold   = bold
        run.italic = italic
    return new_p


def append_sentence(para, sentence):
    """Append text to existing paragraph, matching the first run's font."""
    run = para.add_run(" " + sentence)
    if para.runs and len(para.runs) > 1:
        src = para.runs[0].font
        f   = run.font
        f.name, f.size = src.name, src.size
    return run


def split_paragraph_at(para, split_text):
    """Split para into two at split_text. Returns (original, new_para) or (original, None)."""
    full = para.text
    idx  = full.find(split_text)
    if idx == -1:
        return para, None
    before = full[:idx].rstrip()
    after  = full[idx:]
    for run in para.runs:
        run.text = ""
    if para.runs:
        para.runs[0].text = before
    new_p = insert_paragraph_after(para, after)
    return para, new_p


# ---------------------------------------------------------------------------
doc = Document(IN)

# Print index map for reference
print("=== Index map (first 180 non-blank paras) ===")
count = 0
for i, p in enumerate(doc.paragraphs):
    t = p.text.strip()
    if t:
        print(f"[{i:3d}] {t[:110]}")
        count += 1
        if count >= 180:
            break
print()

# ---------------------------------------------------------------------------
# HIGH-1: Add emerging-market finding preview to Introduction
# ---------------------------------------------------------------------------
# Locate the roadmap paragraph ("Given this background, this paper will examine")
# and insert a new paragraph immediately after it.
idx = find_para_idx(doc, "Given this background, this paper will examine")
if idx is not None:
    new_text = (
        "Our main findings are as follows. Fixed exchange rate arrangements "
        "reduce the probability of explosive current account dynamics by "
        "approximately nine percentage points; manufacturing export "
        "specialisation operates similarly. Crucially, the model fits "
        "substantially better for emerging market economies "
        "(McFadden R\u00b2 = 0.19 versus 0.05 for the full sample), "
        "indicating that exchange rate and trade policy are most consequential "
        "for middle-income countries, where institutional constraints on "
        "external adjustment are less deeply embedded than in advanced economies."
    )
    insert_paragraph_after(doc.paragraphs[idx], new_text)
    print(f"HIGH-1 done: inserted findings preview after para [{idx}]")
else:
    print("HIGH-1 WARNING: roadmap paragraph not found")


# ---------------------------------------------------------------------------
# HIGH-2: Add 77->74 country-drop explanation
# ---------------------------------------------------------------------------
# Locate the sentence "The complete sample covers 74 countries" in Section 5.1.2
idx = find_para_idx(doc, "The complete sample covers 74 countries", start=120, end=160)
if idx is not None:
    append_sentence(
        doc.paragraphs[idx],
        "Three countries included in the unit root analysis (Section 4) "
        "are excluded from the MNL estimation because they have no usable "
        "observations after listwise deletion on covariates, reducing the "
        "sample from 77 to 74 countries."
    )
    print(f"HIGH-2 done: appended 77\u219274 note to para [{idx}]")
else:
    print("HIGH-2 WARNING: '74 countries' paragraph not found in range 120-160")


# ---------------------------------------------------------------------------
# MED-3: Expand post-GFC institutional discussion
# ---------------------------------------------------------------------------
# Locate the post-GFC results paragraph (Section 5.5.2)
idx = find_para_idx(doc, "Model fit more than doubles in the post-GFC period", start=160)
if idx is not None:
    para = doc.paragraphs[idx]
    old  = ("consistent with heightened international scrutiny of global "
            "imbalances and increased cross-country differentiation by "
            "exchange rate arrangement and export structure")
    new  = ("consistent with heightened post-crisis international scrutiny "
            "of global imbalances. Three institutional developments reinforce "
            "this interpretation: the G20 Mutual Assessment Process "
            "(MAP, launched 2009) elevated exchange rate and current account "
            "adjustment to the centre of multilateral policy coordination; "
            "IMF Article IV surveillance intensified its focus on external "
            "sector assessments and currency misalignment from 2012 onward; "
            "and Basel III capital and liquidity requirements (phased from 2013) "
            "increased cross-country differentiation in financial-sector capacity "
            "to sustain external imbalances, sharpening the role of exchange rate "
            "arrangements and export structure as regime determinants")
    full = para.text
    if old in full:
        new_full = full.replace(old, new)
        for run in para.runs:
            run.text = ""
        if para.runs:
            para.runs[0].text = new_full
        print(f"MED-3 done: expanded post-GFC institutional discussion in para [{idx}]")
    else:
        # Append institutional context if exact string not found
        append_sentence(
            para,
            "Three post-crisis institutional developments reinforce this finding: "
            "the G20 Mutual Assessment Process (MAP, 2009) placed exchange rate "
            "and current account adjustment at the centre of multilateral "
            "coordination; IMF Article IV surveillance intensified external sector "
            "assessments from 2012; and Basel III requirements (phased from 2013) "
            "sharpened cross-country differentiation in capacity to sustain "
            "external imbalances."
        )
        print(f"MED-3 done (append): post-GFC para [{idx}]")
else:
    print("MED-3 WARNING: post-GFC paragraph not found")


# ---------------------------------------------------------------------------
# MED-4: Add top explosive-country examples
# ---------------------------------------------------------------------------
# The explosive regime base-rate is described in Section 5.4 (~para 160).
# Find "134 observations" or "5.5%" in that section.
idx = find_para_idx(doc, "134 observations", start=155, end=175)
if idx is None:
    idx = find_para_idx(doc, "5.5% of the sample", start=155, end=175)
if idx is not None:
    append_sentence(
        doc.paragraphs[idx],
        "The explosive regime is geographically concentrated: Israel "
        "(48 country-year observations), Hungary (20), Malaysia (18), "
        "Germany (15), and Armenia (13) together account for the majority "
        "of explosive episodes, spanning both chronic-deficit trajectories "
        "(Hungary, Armenia) and persistent-surplus ones (Malaysia, Germany)."
    )
    print(f"MED-4 done: added country examples to para [{idx}]")
else:
    # Fallback: find in Section 4 discussion of explosive states
    idx2 = find_para_idx(doc, "explosive states are episodic", start=95, end=125)
    if idx2 is None:
        idx2 = find_para_idx(doc, "low frequency of explosive", start=95, end=125)
    if idx2 is not None:
        append_sentence(
            doc.paragraphs[idx2],
            "The explosive regime is geographically concentrated: Israel "
            "(48 country-year observations), Hungary (20), Malaysia (18), "
            "Germany (15), and Armenia (13) together account for the majority "
            "of explosive episodes, spanning both chronic-deficit and "
            "persistent-surplus trajectories."
        )
        print(f"MED-4 done (fallback Section 4): added country examples to para [{idx2}]")
    else:
        print("MED-4 WARNING: explosive distribution paragraph not found")


# ---------------------------------------------------------------------------
# LOW-5: Split two-step / IIA limitations paragraph into two
# ---------------------------------------------------------------------------
idx = find_para_idx(doc, "Two limitations of the two-step framework", start=130, end=150)
if idx is not None:
    para = doc.paragraphs[idx]
    _, new_p = split_paragraph_at(para, "Second, the multinomial logit")
    if new_p is not None:
        print(f"LOW-5 done: split limitations para [{idx}] into two")
    else:
        print(f"LOW-5 WARNING: 'Second, the multinomial logit' not found in para [{idx}]")
else:
    print("LOW-5 WARNING: limitations paragraph not found")


# ---------------------------------------------------------------------------
# LOW-6: Small-cluster caveat for advanced-economy subgroup (Section 5.5.1)
# ---------------------------------------------------------------------------
# Section 5.5.1 is around para [167]-[170].
# Find the paragraph that discusses the advanced-economy interpretation
# ("institutional literature" or "advanced economies possess more").
idx = find_para_idx(doc,
    "The finding is consistent with the broader institutional literature",
    start=160, end=185)
if idx is None:
    idx = find_para_idx(doc, "advanced economies possess", start=160, end=185)
if idx is None:
    # Try the section itself
    idx = find_para_idx(doc, "5.5.1", start=155, end=175)
if idx is not None:
    append_sentence(
        doc.paragraphs[idx],
        "An additional caveat applies to the advanced-economy subgroup: "
        "with only 23 country clusters, cluster-robust standard errors "
        "may be less reliable than in the full sample "
        "(Cameron and Miller 2015 suggest at least 30 clusters for "
        "conventional cluster-robust inference), so these results should "
        "be interpreted with additional caution."
    )
    print(f"LOW-6 done: added small-cluster caveat to para [{idx}]")
else:
    print("LOW-6 WARNING: Section 5.5.1 institutional paragraph not found")


# ---------------------------------------------------------------------------
# LOW-7: MS-model structure homogeneity (Section 3.2 / 3.2.1)
# ---------------------------------------------------------------------------
# Find the formal model specification paragraph that describes the MS-AR model.
# Look for "Formally" or "unobserved regime indicators" in Section 3.2 range.
idx = find_para_idx(doc, "Formally, let", start=65, end=95)
if idx is None:
    idx = find_para_idx(doc, "unobserved regime indicators", start=65, end=95)
if idx is not None:
    append_sentence(
        doc.paragraphs[idx],
        "We impose a common two-regime AR(1) structure across all countries "
        "to ensure comparability of regime classifications; allowing "
        "country-specific lag orders would complicate cross-country aggregation "
        "and is left for future work."
    )
    print(f"LOW-7 done: added MS homogeneity note to para [{idx}]")
else:
    print("LOW-7 WARNING: MS formal model paragraph not found in range 65-95")


# ---------------------------------------------------------------------------
# Save
# ---------------------------------------------------------------------------
doc.save(OUT)
print(f"\nSaved: {OUT}")
print("=== All edits applied ===")
