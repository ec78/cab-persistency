"""
apply_pass4_edits_v6.py
=======================
Pass 4 edits: working_draft_ec_2025_v5_Mar2026.docx → working_draft_ec_2025_v6_Mar2026.docx

Addresses three remaining gaps identified in audit:
  F1  Add explosive recall % (3 percent) to confusion matrix paragraph (David)
  F2  Add fdgap discussion + BIC/sample comparability caveat in Section 5.3.1 (Rachel)
  F3  Compress Section 3.1 from 11 paragraphs to 2 substantive paragraphs (James)

Run from repo root. Requires: python-docx.
"""

import copy
import sys
from docx import Document
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

INPUT  = r"c:\Users\eclow\Documents\GitHub\cab-persistency\updated_2025\working_draft_ec_2025_v5_Mar2026.docx"
OUTPUT = r"c:\Users\eclow\Documents\GitHub\cab-persistency\updated_2025\working_draft_ec_2025_v6_Mar2026.docx"


# ---------------------------------------------------------------------------
# Helpers
# ---------------------------------------------------------------------------

def replace_in_para(para, old, new):
    full = para.text
    if old not in full:
        return False
    if para.runs:
        para.runs[0].text = full.replace(old, new)
        for r in para.runs[1:]:
            r.text = ""
    return True


def insert_paragraph_after(ref_para, text):
    new_para = OxmlElement('w:p')
    ref_pPr = ref_para._element.find(qn('w:pPr'))
    if ref_pPr is not None:
        new_para.append(copy.deepcopy(ref_pPr))
    new_r = OxmlElement('w:r')
    if ref_para.runs:
        ref_rPr = ref_para.runs[0]._element.find(qn('w:rPr'))
        if ref_rPr is not None:
            new_r.append(copy.deepcopy(ref_rPr))
    new_t = OxmlElement('w:t')
    new_t.text = text
    new_t.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')
    new_r.append(new_t)
    new_para.append(new_r)
    ref_para._element.addnext(new_para)
    return new_para


def remove_paragraph(para):
    """Remove a paragraph from the document."""
    p = para._element
    parent = p.getparent()
    if parent is not None:
        parent.remove(p)


def set_para_text(para, text):
    """Replace all content in a paragraph with new text (preserving first run's style)."""
    if para.runs:
        para.runs[0].text = text
        for r in para.runs[1:]:
            r.text = ""
    else:
        run = para.add_run(text)


# ---------------------------------------------------------------------------
# Load document
# ---------------------------------------------------------------------------

doc = Document(INPUT)
paras = [p for p in doc.paragraphs if p.text.strip()]
print(f"Loaded. Non-empty paragraphs: {len(paras)}")


# ---------------------------------------------------------------------------
# F1: Add explicit 3% explosive recall to confusion matrix paragraph
# ---------------------------------------------------------------------------

for p in doc.paragraphs:
    if 'Table C1 in the Appendix reports a confusion matrix' in p.text:
        old = ("The model correctly classifies unit root observations with high recall "
               "but has limited ability to identify explosive episodes.")
        new = ("The model correctly classifies unit root observations with high recall "
               "but has very limited ability to identify explosive episodes: explosive "
               "recall is approximately 3 percent, meaning the model correctly identifies "
               "only about 4 of the 134 explosive country-years in-sample.")
        result = replace_in_para(p, old, new)
        print(f"F1 (explosive recall %): {'done' if result else 'FAILED'}")
        if not result:
            print("  Text not matched. Para snippet:", p.text[:300])
        break


# ---------------------------------------------------------------------------
# F2: Add fdgap discussion + BIC comparability caveat after crisis dummy para
# ---------------------------------------------------------------------------

CRISIS_PARA_PREFIX = "As a direct test of this conclusion, we include indicator variables for the Global Financial Crisis"
crisis_para = None
for p in doc.paragraphs:
    if p.text.strip().startswith(CRISIS_PARA_PREFIX):
        crisis_para = p
        break

if crisis_para:
    FDGAP_TEXT = (
        "The financial development gap specification (C00342_fdgap) yields a modest improvement "
        "in McFadden R\u00b2 (0.067 versus 0.050) but reduces the estimation sample by "
        "approximately 16 percent due to missing financial development data (n\u202f=\u202f2,064 "
        "versus n\u202f=\u202f2,451 for the baseline). Given the cost in sample coverage relative "
        "to the marginal gain in fit, this variable is excluded from the preferred specification. "
        "A comparability caveat applies to BIC/obs values across specifications with different "
        "estimation samples: because BIC/obs is computed on distinct observation sets, the values "
        "for C00342_fdgap and C00058 are not strictly comparable. Readers should interpret the "
        "cross-specification BIC comparison as indicative rather than definitive for models "
        "estimated on different samples."
    )
    insert_paragraph_after(crisis_para, FDGAP_TEXT)
    print("F2 (fdgap + BIC caveat): inserted after crisis dummy paragraph")
else:
    print("F2: crisis paragraph not found")


# ---------------------------------------------------------------------------
# F3: Compress Section 3.1 from 11 paragraphs to 2 substantive paragraphs
#
# Current structure (non-empty paras in Section 3.1):
#   [0] Heading "3.1. Linear unit root tests"          <- KEEP heading text
#   [1] "We begin by revisiting..."                    <- REPLACE with merged Para 1
#   [2] "As emphasized in Section 2..."                <- REMOVE
#   [3] "For this reason, we interpret..."             <- REMOVE
#   [4] "3.1.2 Unit Root Testing Results"              <- REMOVE (subsection heading)
#   [5] "Appendix 2 reports standard linear..."        <- REPLACE with merged Para 2
#   [6] "In many cases, Phillips-Perron..."            <- REMOVE
#   [7] "Cross-country differences are pronounced..."  <- REMOVE
#   [8] "" (blank)                                     <- already filtered out
#   [9] "Figure 4. Linear Unit Root Results"           <- KEEP (figure caption)
#  [10] "Figure 4 summarizes how standard..."          <- REMOVE (merge into Para 2)
#  [11] "Rather than interpreting these outcomes..."   <- REMOVE (merge into Para 2)
# ---------------------------------------------------------------------------

MERGED_PARA1 = (
    "Standard linear unit root tests provide a natural benchmark for assessing current account "
    "persistence. We apply ADF, DFGLS, and Phillips\u2013Perron tests to current account "
    "balances expressed as a share of GDP across our 77-country sample (full results in "
    "Appendix 2). These tests assume that the mean, variance, and persistence of the current "
    "account process are time-invariant \u2014 an assumption that may be too restrictive when "
    "dynamics shift over time due to structural change, policy regime shifts, or episodic "
    "adjustment. We therefore treat their results as a descriptive baseline rather than a "
    "definitive assessment, using them to motivate the regime-switching framework in Section 3.2."
)

MERGED_PARA2 = (
    "Figure 4 summarizes the classification outcomes across tests. The central message is that "
    "conclusions are highly test-dependent. The ADF and DFGLS tests are relatively conservative, "
    "with roughly half of all series failing to reject a unit root at conventional levels; "
    "Phillips\u2013Perron statistics reject the null for a substantially larger share of "
    "countries, consistent with those tests' greater sensitivity to time variation in persistence "
    "and volatility. The KPSS test rejects stationarity for more than 80 percent of series. "
    "Cross-country heterogeneity is also pronounced: some economies show consistent mean-reverting "
    "evidence across multiple tests while others produce contradictory results, likely reflecting "
    "heterogeneity in policy frameworks, shock exposure, and available sample length. Rather than "
    "treating this divergence as evidence for or against stationarity, we interpret it as "
    "confirmation that a single constant-parameter linear framework is inadequate and that time "
    "variation in current account dynamics must be accommodated \u2014 motivating the "
    "Markov-switching approach in Section 3.2."
)

# Collect all paragraphs in Section 3.1 (including empty ones, for removal)
sec31_paras = []
in_31 = False
for p in doc.paragraphs:
    t = p.text.strip()
    if t == '3.1. Linear unit root tests':
        in_31 = True
        sec31_paras.append(p)
        continue
    if in_31:
        if t.startswith('3.2'):
            break
        sec31_paras.append(p)

print(f"F3: found {len(sec31_paras)} paragraphs in Section 3.1 (including empty and heading)")

if len(sec31_paras) >= 5:
    # Identify key paragraphs by content
    heading_para = None
    intro_para = None     # "We begin by revisiting..."
    sub_heading_para = None   # "3.1.2"
    results_para = None   # "Appendix 2 reports..."
    fig4_caption = None   # "Figure 4. Linear Unit Root Results"
    fig4_discuss = None   # "Figure 4 summarizes how..."
    conclusion_para = None  # "Rather than interpreting..."
    to_remove = []

    for p in sec31_paras:
        t = p.text.strip()
        if t == '3.1. Linear unit root tests':
            heading_para = p
        elif t.startswith('We begin by revisiting'):
            intro_para = p
        elif t.startswith('As emphasized in Section 2'):
            to_remove.append(p)
        elif t.startswith('For this reason, we interpret'):
            to_remove.append(p)
        elif '3.1.2' in t:
            sub_heading_para = p
        elif t.startswith('Appendix 2 reports standard'):
            results_para = p
        elif t.startswith('In many cases, Phillips'):
            to_remove.append(p)
        elif t.startswith('Cross-country differences are pronounced'):
            to_remove.append(p)
        elif t == 'Figure 4. Linear Unit Root Results':
            fig4_caption = p
        elif t.startswith('Figure 4 summarizes how standard'):
            fig4_discuss = p
        elif t.startswith('Rather than interpreting these outcomes'):
            conclusion_para = p
        elif t == '':
            to_remove.append(p)

    # Replace intro_para with MERGED_PARA1
    if intro_para:
        set_para_text(intro_para, MERGED_PARA1)
        print("  Set intro paragraph to merged para 1")
    else:
        print("  WARNING: intro_para not found")

    # Replace results_para with MERGED_PARA2
    if results_para:
        set_para_text(results_para, MERGED_PARA2)
        print("  Set results paragraph to merged para 2")
    else:
        print("  WARNING: results_para not found")

    # Remove paragraphs that are now merged or obsolete
    removals = [p for p in to_remove]
    if sub_heading_para:
        removals.append(sub_heading_para)
    if fig4_discuss:
        removals.append(fig4_discuss)
    if conclusion_para:
        removals.append(conclusion_para)

    for p in removals:
        remove_paragraph(p)
        print(f"  Removed: {repr(p.text.strip()[:60])}")

    print(f"F3 (Section 3.1 compression): done — removed {len(removals)} paragraphs")
else:
    print("F3: not enough paragraphs found in Section 3.1; skipping")


# ---------------------------------------------------------------------------
# Save
# ---------------------------------------------------------------------------

doc.save(OUTPUT)
print(f"\nSaved: {OUTPUT}")
print("Pass 4 complete.")
