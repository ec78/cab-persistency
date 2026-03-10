"""
apply_pass2_edits.py
Implements Pass 2 open items for working_draft_ec_2025_v3_Mar2026.docx.
Saves output as working_draft_ec_2025_v4_Mar2026.docx.

Changes implemented:
  HIGH-1 : Add emerging-market finding preview to Introduction
  HIGH-2 : Add 77->74 country-drop explanation in Section 5.1.2
  MED-3  : Expand post-GFC institutional discussion (Section 5.5.2)
  MED-4  : Add top-explosive-country examples in baseline-results discussion
  LOW-5  : Split two-step / IIA limitations into two paragraphs
  LOW-6  : Add small-cluster caveat for advanced-economy subgroup (Section 5.5.1)
  LOW-7  : Add MS-model structure homogeneity statement (Section 3.2)
"""

import sys, io
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding="utf-8", errors="replace")

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt
from lxml import etree
import copy

IN  = (r"c:/Users/eclow/Documents/GitHub/cab-persistency/"
       r"updated_2025/working_draft_ec_2025_v3_Mar2026.docx")
OUT = (r"c:/Users/eclow/Documents/GitHub/cab-persistency/"
       r"updated_2025/working_draft_ec_2025_v4_Mar2026.docx")


# ---------------------------------------------------------------------------
# Helpers
# ---------------------------------------------------------------------------

def find_para_idx(doc, search_text, case_sensitive=False):
    """Return the first paragraph index whose text contains search_text."""
    needle = search_text if case_sensitive else search_text.lower()
    for i, p in enumerate(doc.paragraphs):
        hay = p.text if case_sensitive else p.text.lower()
        if needle in hay:
            return i
    return None


def get_para_style(para):
    """Return the style name of a paragraph (safe)."""
    try:
        return para.style.name
    except Exception:
        return "Normal"


def clone_run_format(source_run, target_run):
    """Copy font properties from source_run to target_run."""
    src_font = source_run.font
    tgt_font = target_run.font
    tgt_font.name      = src_font.name
    tgt_font.size      = src_font.size
    tgt_font.bold      = src_font.bold
    tgt_font.italic    = src_font.italic
    tgt_font.underline = src_font.underline


def insert_paragraph_after(ref_para, text, style_name=None, italic=False, bold=False):
    """
    Insert a new paragraph immediately after ref_para.
    Returns the new Paragraph object.
    """
    from docx.text.paragraph import Paragraph as DocxPara

    # Build new <w:p> element and insert after ref
    new_p_elem = copy.deepcopy(ref_para._element)
    # Clear all runs from the clone
    for r in new_p_elem.findall(qn("w:r")):
        new_p_elem.remove(r)
    ref_para._element.addnext(new_p_elem)

    # Wrap in Paragraph object
    new_para = DocxPara(new_p_elem, ref_para._parent)

    if style_name:
        try:
            new_para.style = ref_para.part.document.styles[style_name]
        except Exception:
            pass

    run = new_para.add_run(text)
    run.bold   = bold
    run.italic = italic

    # Match font from first run of ref_para, if any
    ref_runs = ref_para.runs
    if ref_runs:
        clone_run_format(ref_runs[0], run)

    return new_para


def append_sentence(para, sentence):
    """Append a sentence (with a leading space) to an existing paragraph."""
    run = para.add_run(" " + sentence)
    runs = para.runs
    if len(runs) > 1:
        clone_run_format(runs[0], run)
    return run


def split_paragraph_at(para, split_text):
    """
    Split para into two at the first occurrence of split_text.
    split_text and everything after it become a new paragraph inserted after para.
    Returns (original_para, new_para).
    """
    full_text = para.text
    idx = full_text.find(split_text)
    if idx == -1:
        print(f"  WARNING: split_text not found: {split_text[:60]!r}")
        return para, None

    before = full_text[:idx].rstrip()
    after  = full_text[idx:]

    # Replace the paragraph's text with 'before'
    # Clear all runs
    for run in para.runs:
        run.text = ""
    para.runs[0].text = before if before else ""

    # Insert new paragraph after with 'after'
    new_para = insert_paragraph_after(para, after)
    return para, new_para


# ---------------------------------------------------------------------------
# Load document
# ---------------------------------------------------------------------------
doc = Document(IN)

# Print paragraph index map for verification
print("=== Paragraph Index Map ===")
for i, p in enumerate(doc.paragraphs):
    txt = p.text.strip()
    if txt:
        print(f"[{i:3d}] {txt[:100]}")

print("\n=== Applying edits ===\n")


# ===========================================================================
# HIGH-1: Add emerging-market finding preview to Introduction
# ---------------------------------------------------------------------------
# Strategy: Find the roadmap paragraph ("Given this background, this paper will
# examine") and insert a new "main findings" paragraph immediately after it.
# ===========================================================================
idx = find_para_idx(doc, "Given this background, this paper will examine")
if idx is not None:
    ref = doc.paragraphs[idx]
    new_text = (
        "The paper's main empirical findings can be summarised as follows. "
        "Fixed exchange rate arrangements reduce the probability of explosive "
        "current account dynamics by approximately nine percentage points; "
        "manufacturing export specialisation operates similarly. "
        "The model fits substantially better for emerging market economies "
        "(McFadden R\u00b2 = 0.19 versus 0.05 for the full sample), indicating "
        "that exchange rate and trade policy are most consequential for middle-income "
        "countries, where policy constraints on external adjustment are less "
        "institutionally embedded than in advanced economies."
    )
    insert_paragraph_after(ref, new_text)
    print(f"HIGH-1 done: inserted findings preview after para [{idx}]")
else:
    print("HIGH-1 WARNING: roadmap paragraph not found")


# ===========================================================================
# HIGH-2: Add 77->74 country-drop explanation in Section 5.1.2
# ---------------------------------------------------------------------------
# Find the paragraph that says "The complete sample covers 74 countries"
# and append the explanation.
# ===========================================================================
idx = find_para_idx(doc, "The complete sample covers 74 countries")
if idx is not None:
    para = doc.paragraphs[idx]
    append_sentence(
        para,
        "Three countries present in the unit root analysis (Section 4) are "
        "excluded from the MNL estimation because missing covariate data leave "
        "them with no usable observations after listwise deletion, reducing the "
        "analysis sample from 77 to 74 countries."
    )
    print(f"HIGH-2 done: appended 77->74 explanation to para [{idx}]")
else:
    print("HIGH-2 WARNING: '74 countries' paragraph not found")


# ===========================================================================
# MED-3: Expand post-GFC institutional discussion (Section 5.5.2)
# ---------------------------------------------------------------------------
# Find the post-GFC paragraph and expand it with G20/IMF/Basel III anchors.
# ===========================================================================
idx = find_para_idx(doc, "Model fit more than doubles in the post-GFC period")
if idx is not None:
    para = doc.paragraphs[idx]
    # Replace the institutional-anchors clause with a richer version
    old_clause = (
        "consistent with heightened international scrutiny of global imbalances "
        "and increased cross-country differentiation by exchange rate arrangement "
        "and export structure"
    )
    new_clause = (
        "consistent with heightened international scrutiny of global imbalances "
        "after the crisis. Institutional developments reinforce this interpretation: "
        "the G20 Mutual Assessment Process (MAP, launched 2009) placed exchange "
        "rate and current account adjustment at the centre of multilateral policy "
        "coordination; IMF Article IV surveillance intensified its focus on external "
        "sector assessments and currency misalignment from 2012 onward; and the "
        "Basel III capital and liquidity requirements (phased in from 2013) "
        "increased cross-country differentiation in financial-sector capacity to "
        "sustain external imbalances. Together, these post-crisis institutional "
        "changes sharpened the policy relevance of exchange rate arrangements and "
        "export structure as regime determinants"
    )
    full = para.text
    if old_clause in full:
        # Rebuild runs: clear and rewrite
        new_full = full.replace(old_clause, new_clause)
        for run in para.runs:
            run.text = ""
        para.runs[0].text = new_full
        print(f"MED-3 done: expanded post-GFC para [{idx}]")
    else:
        # Clause not found verbatim — append instead
        append_sentence(
            para,
            "Post-crisis institutional developments sharpen this interpretation: "
            "the G20 Mutual Assessment Process (MAP, 2009) placed exchange rate "
            "and current account adjustment at the centre of multilateral policy "
            "coordination; IMF Article IV surveillance intensified its focus on "
            "external sector assessments from 2012 onward; and Basel III capital "
            "requirements (phased from 2013) increased cross-country differentiation "
            "in financial-sector capacity to sustain external imbalances."
        )
        print(f"MED-3 done (append fallback): expanded post-GFC para [{idx}]")
else:
    print("MED-3 WARNING: post-GFC paragraph not found")


# ===========================================================================
# MED-4: Add top-explosive-country examples
# ---------------------------------------------------------------------------
# Find the paragraph that first describes the outcome distribution
# (mentions "5.5%" or "134" explosive observations) and add a sentence
# listing the most frequently classified countries.
# ===========================================================================
idx = find_para_idx(doc, "5.5%")
if idx is None:
    idx = find_para_idx(doc, "explosive")   # fallback: first mention of regime distribution
if idx is not None:
    # Walk forward a few paragraphs to find one that discusses the sample composition
    for offset in range(0, 6):
        candidate = doc.paragraphs[idx + offset]
        t = candidate.text.lower()
        if "explosive" in t and ("38" in t or "56" in t or "5.5" in t or "134" in t):
            append_sentence(
                candidate,
                "The explosive regime is geographically concentrated: Israel "
                "(48 country-year observations), Hungary (20), Malaysia (18), "
                "Germany (15), and Armenia (13) account for the majority of "
                "explosive episodes, spanning both chronic-deficit and persistent-surplus "
                "trajectories."
            )
            print(f"MED-4 done: added country examples to para [{idx + offset}]")
            break
    else:
        # Fallback: append to the paragraph found
        append_sentence(
            doc.paragraphs[idx],
            "The explosive regime is geographically concentrated: Israel "
            "(48 country-year observations), Hungary (20), Malaysia (18), "
            "Germany (15), and Armenia (13) account for the majority of "
            "explosive episodes, spanning both chronic-deficit and persistent-surplus "
            "trajectories."
        )
        print(f"MED-4 done (fallback): added country examples to para [{idx}]")
else:
    print("MED-4 WARNING: explosive distribution paragraph not found")


# ===========================================================================
# LOW-5: Split two-step / IIA limitations into two paragraphs
# ---------------------------------------------------------------------------
# Find the combined limitations paragraph starting with
# "Two limitations of the two-step framework" and split on "Second,"
# ===========================================================================
idx = find_para_idx(doc, "Two limitations of the two-step framework")
if idx is not None:
    para = doc.paragraphs[idx]
    before, new_para = split_paragraph_at(para, "Second, the multinomial logit")
    if new_para is not None:
        print(f"LOW-5 done: split limitations para [{idx}] into two paragraphs")
    else:
        print(f"LOW-5 WARNING: split point not found in para [{idx}]")
else:
    print("LOW-5 WARNING: limitations paragraph not found")


# ===========================================================================
# LOW-6: Small-cluster caveat for advanced-economy subgroup (Section 5.5.1)
# ---------------------------------------------------------------------------
# Find the paragraph discussing the advanced-economy subgroup fit issue
# (mentions "15" explosive observations or "23 clusters") and append caveat.
# ===========================================================================
idx = find_para_idx(doc, "23 clusters")
if idx is None:
    idx = find_para_idx(doc, "advanced economy")
if idx is None:
    idx = find_para_idx(doc, "Advanced economies")
if idx is not None:
    # Walk up to 5 paragraphs to find one that discusses the subgroup limitation
    for offset in range(0, 8):
        candidate = doc.paragraphs[idx + offset]
        t = candidate.text.lower()
        if ("advanced" in t and ("15" in t or "23" in t or "reliable" in t or "few explosive" in t)):
            append_sentence(
                candidate,
                "An additional caveat applies: with only 23 country clusters in the "
                "advanced-economy subsample, cluster-robust standard errors may be "
                "less reliable than in the full sample (Cameron and Miller 2015 "
                "recommend at least 30 clusters for conventional cluster-robust "
                "inference), so the advanced-economy results should be interpreted "
                "with appropriate caution."
            )
            print(f"LOW-6 done: added small-cluster caveat to para [{idx + offset}]")
            break
    else:
        # Append to the first matching paragraph
        append_sentence(
            doc.paragraphs[idx],
            "An additional caveat applies: with only 23 country clusters in the "
            "advanced-economy subsample, cluster-robust standard errors may be "
            "less reliable than in the full sample (Cameron and Miller 2015 "
            "recommend at least 30 clusters for conventional cluster-robust "
            "inference), so the advanced-economy results should be interpreted "
            "with appropriate caution."
        )
        print(f"LOW-6 done (fallback): added small-cluster caveat to para [{idx}]")
else:
    print("LOW-6 WARNING: advanced-economy paragraph not found")


# ===========================================================================
# LOW-7: Add MS-model structure homogeneity statement (Section 3.2)
# ---------------------------------------------------------------------------
# Find the MS model description paragraph and append a note about
# imposing the same switching structure across all countries.
# ===========================================================================
idx = find_para_idx(doc, "Markov-switching")
# Walk forward to find the first paragraph in Section 3.2 context
ms_para_idx = None
for offset in range(0, 60):
    candidate = doc.paragraphs[idx + offset]
    t = candidate.text.lower()
    if ("markov" in t and "switching" in t and
            ("autoregressive" in t or "ar(" in t or "transition" in t or "regime" in t)):
        ms_para_idx = idx + offset
        break

if ms_para_idx is not None:
    para = doc.paragraphs[ms_para_idx]
    # Only add if homogeneity sentence not already present
    if "same switching structure" not in para.text.lower() and "homogeneous" not in para.text.lower():
        append_sentence(
            para,
            "We impose a homogeneous switching structure across all countries "
            "(two regimes, AR(1) within each regime) to ensure comparability of "
            "regime classifications; allowing heterogeneous lag orders would "
            "complicate cross-country aggregation and is left for future work."
        )
        print(f"LOW-7 done: added MS homogeneity note to para [{ms_para_idx}]")
    else:
        print(f"LOW-7 skipped: homogeneity already noted in para [{ms_para_idx}]")
else:
    print("LOW-7 WARNING: MS autoregressive paragraph not found")


# ===========================================================================
# Save
# ===========================================================================
doc.save(OUT)
print(f"\nSaved: {OUT}")
print("=== All edits complete ===")
